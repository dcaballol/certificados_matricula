"""
Generador de Certificados de Matrícula
SLEP Santa Corina
VERSIÓN MEJORADA - Funciona con cualquier template
"""

from docx import Document
from datetime import datetime
import io
import re


class GeneradorCertificado:
    """Clase para generar certificados de matrícula personalizados"""
    
    def __init__(self, template_path):
        """
        Inicializa el generador con la ruta del template
        
        Args:
            template_path (str): Ruta al archivo .docx template
        """
        self.template_path = template_path
    
    def generar_certificado(self, datos_estudiante, fecha_emision=None):
        """
        Genera un certificado de matrícula personalizado
        
        Args:
            datos_estudiante (dict): Diccionario con los datos del estudiante
                - nombre: Nombre completo
                - run: RUN formateado
                - establecimiento: Nombre del establecimiento
                - rbd: Código RBD
                - curso: Curso completo (ej: "6° básico C")
                - año: Año escolar
            fecha_emision (datetime, optional): Fecha de emisión del certificado
            
        Returns:
            io.BytesIO: Documento Word en memoria
        """
        # Cargar el template
        doc = Document(self.template_path)
        
        # Usar fecha actual si no se proporciona
        if fecha_emision is None:
            fecha_emision = datetime.now()
        
        # Formatear fecha
        fecha_formateada = self._formatear_fecha(fecha_emision)
        
        # Reemplazar en párrafos
        for para in doc.paragraphs:
            self._reemplazar_en_texto(para, datos_estudiante, fecha_formateada)
        
        # Reemplazar en tablas (si las hay)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._reemplazar_en_texto(para, datos_estudiante, fecha_formateada)
        
        # Guardar en memoria
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def _reemplazar_en_texto(self, para, datos, fecha):
        """
        Reemplaza los placeholders en un párrafo usando patrones inteligentes
        
        Args:
            para: Párrafo de python-docx
            datos (dict): Datos del estudiante
            fecha (str): Fecha formateada
        """
        # Obtener el texto completo del párrafo
        texto_original = para.text
        
        if not texto_original.strip():
            return
        
        texto_nuevo = texto_original
        
        # PATRÓN 1: Buscar "Don(a) NOMBRE, RUN"
        # Captura nombres en mayúsculas antes de una coma
        patron_nombre = r'Don\(a\)\s+([A-ZÁÉÍÓÚÑ\s]+?)(?=,)'
        if re.search(patron_nombre, texto_nuevo, re.IGNORECASE):
            texto_nuevo = re.sub(
                patron_nombre, 
                f"Don(a) {datos.get('nombre', '').upper()}", 
                texto_nuevo,
                flags=re.IGNORECASE
            )
        
        # PATRÓN 2: Buscar RUN con formato XX.XXX.XXX-X
        patron_run = r'\d{1,2}\.\d{3}\.\d{3}-[\dKk]'
        if re.search(patron_run, texto_nuevo) and datos.get('run'):
            texto_nuevo = re.sub(patron_run, datos.get('run', ''), texto_nuevo)
        
        # PATRÓN 3: Buscar "RBD" seguido de números
        patron_rbd = r'RBD\s+\d+'
        if re.search(patron_rbd, texto_nuevo) and datos.get('rbd'):
            texto_nuevo = re.sub(patron_rbd, f"RBD  {datos.get('rbd', '')}", texto_nuevo)
        
        # PATRÓN 4: Buscar nombre de curso (X° básico/medio)
        patron_curso = r'\d+°\s+(básico|medio)\s+[A-Z]'
        if re.search(patron_curso, texto_nuevo, re.IGNORECASE) and datos.get('curso'):
            texto_nuevo = re.sub(
                patron_curso,
                datos.get('curso', ''),
                texto_nuevo,
                flags=re.IGNORECASE
            )
        
        # PATRÓN 5: Buscar año (4 dígitos consecutivos)
        patron_anio = r'\b202\d\b'
        if re.search(patron_anio, texto_nuevo) and datos.get('año'):
            texto_nuevo = re.sub(patron_anio, str(datos.get('año', '')), texto_nuevo)
        
        # PATRÓN 6: Buscar fecha completa "DD de MES del YYYY"
        patron_fecha = r'\d{1,2}\s+de\s+\w+\s+del\s+\d{4}'
        if re.search(patron_fecha, texto_nuevo):
            texto_nuevo = re.sub(patron_fecha, fecha, texto_nuevo)
        
        # PATRÓN 7: Buscar nombre de establecimiento (todo en mayúsculas)
        # Reemplazar nombres largos en mayúsculas (más de 3 palabras)
        patron_establecimiento = r'[A-ZÁÉÍÓÚÑ]+(?:\s+[A-ZÁÉÍÓÚÑ]+){2,}'
        matches = re.findall(patron_establecimiento, texto_nuevo)
        if matches and datos.get('establecimiento'):
            # Reemplazar el match más largo (probablemente el nombre del establecimiento)
            nombre_mas_largo = max(matches, key=len)
            if len(nombre_mas_largo) > 10:  # Solo si es suficientemente largo
                texto_nuevo = texto_nuevo.replace(
                    nombre_mas_largo,
                    datos.get('establecimiento', '').upper()
                )
        
        # Si hubo cambios, actualizar el párrafo
        if texto_nuevo != texto_original:
            # Limpiar runs existentes
            for run in para.runs:
                run.text = ''
            
            # Agregar el nuevo texto
            if para.runs:
                para.runs[0].text = texto_nuevo
            else:
                para.add_run(texto_nuevo)
    
    def _formatear_fecha(self, fecha):
        """
        Formatea una fecha para el certificado
        
        Args:
            fecha (datetime): Fecha a formatear
            
        Returns:
            str: Fecha formateada (ej: "20 de enero del 2026")
        """
        meses = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }
        
        dia = fecha.day
        mes = meses[fecha.month]
        año = fecha.year
        
        return f"{dia} de {mes} del {año}"
    
    @staticmethod
    def preparar_datos_estudiante(row):
        """
        Prepara los datos de un estudiante desde un DataFrame row
        
        Args:
            row: Fila de pandas DataFrame
            
        Returns:
            dict: Datos formateados para el certificado
        """
        from utils import formatear_run, formatear_curso
        
        # Formatear RUN con DV
        run_formateado = formatear_run(row['SAL_RUN'])
        
        # Formatear curso
        curso = formatear_curso(
            row['COD_GRADO_GLOSA_PRE'],
            row['LET_CUR_PRE']
        )
        
        return {
            'nombre': row.get('NOMBRE_ESTUDIANTE', 'NOMBRE NO DISPONIBLE'),
            'run': run_formateado,
            'establecimiento': row['NOM_RBD'],
            'rbd': row['RBD_PRE'],
            'curso': curso,
            'año': row['ANO_ESCOLAR']
        }
